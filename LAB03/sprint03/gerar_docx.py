"""
Converte relatorio.md para relatorio.docx usando python-docx.
Suporta: títulos, parágrafos, negrito/itálico/code inline,
tabelas, listas, imagens e blocos de código/mermaid.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SPRINT03_DIR = Path(__file__).resolve().parent
MD_PATH  = SPRINT03_DIR / "relatorio.md"
OUT_PATH = SPRINT03_DIR / "relatorio.docx"


# ── Helpers de formatação inline ──────────────────────────────────────────────

INLINE_PATTERNS = [
    # order matters: bold-italic before bold before italic
    (re.compile(r"\*\*\*(.+?)\*\*\*"), "bold_italic"),
    (re.compile(r"\*\*(.+?)\*\*"),      "bold"),
    (re.compile(r"\*(.+?)\*"),          "italic"),
    (re.compile(r"_(.+?)_"),            "italic"),
    (re.compile(r"`(.+?)`"),            "code"),
]

def add_inline(run_parent, text: str) -> None:
    """Parse inline markdown and add formatted runs to a paragraph."""
    # Split text by inline patterns, preserving matched groups
    tokens = _tokenize(text)
    for kind, value in tokens:
        run = run_parent.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "bold_italic":
            run.bold = True
            run.italic = True
        elif kind == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


def _tokenize(text: str):
    """Return list of (kind, value) tuples for inline markdown."""
    combined = re.compile(
        r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_|`(.+?)`)"
    )
    result = []
    last = 0
    for m in combined.finditer(text):
        if m.start() > last:
            result.append(("text", text[last:m.start()]))
        raw = m.group(0)
        if raw.startswith("***"):
            result.append(("bold_italic", m.group(2)))
        elif raw.startswith("**"):
            result.append(("bold", m.group(3)))
        elif raw.startswith("`"):
            result.append(("code", m.group(6)))
        else:
            result.append(("italic", m.group(4) or m.group(5)))
        last = m.end()
    if last < len(text):
        result.append(("text", text[last:]))
    return result


# ── Helpers de estilo de documento ───────────────────────────────────────────

def set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for level, size, bold in [
        ("Heading 1", 20, True),
        ("Heading 2", 16, True),
        ("Heading 3", 13, True),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Code block style
    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", 1)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.4)


def add_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "AAAAAA")
        borders.append(el)
    tblPr.append(borders)


# ── Parser principal ───────────────────────────────────────────────────────────

def parse_md(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ── Blank line
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            doc.add_paragraph().add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # ── Headings
        hm = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2).strip()
            # Strip anchor links like {#...}
            text = re.sub(r"\s*\{#[^}]+\}", "", text)
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_paragraph(style=heading_map.get(level, "Heading 3"))
            add_inline(p, text)
            i += 1
            continue

        # ── Fenced code block (``` or mermaid)
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # closing ```
            if lang == "mermaid":
                p = doc.add_paragraph()
                p.style = "Normal"
                r = p.add_run("[Diagrama de fluxo - ver versão online do relatório]")
                r.italic = True
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            else:
                for cl in code_lines:
                    p = doc.add_paragraph(style="Code Block")
                    p.add_run(cl)
            continue

        # ── Image
        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_m:
            alt  = img_m.group(1)
            path = SPRINT03_DIR / img_m.group(2)
            if path.exists():
                try:
                    doc.add_picture(str(path), width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"[Imagem: {alt}]").italic = True
            else:
                doc.add_paragraph(f"[Imagem não encontrada: {alt}]").italic = True
            # Caption (next line if it starts with *)
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("*"):
                cap_text = lines[i + 1].strip().strip("*")
                cap = doc.add_paragraph(cap_text)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                i += 2
            else:
                i += 1
            continue

        # ── Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # ── Unordered list
        ul_m = re.match(r"^[-*+]\s+(.*)", stripped)
        if ul_m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, ul_m.group(1))
            i += 1
            continue

        # ── Ordered list
        ol_m = re.match(r"^\d+\.\s+(.*)", stripped)
        if ol_m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, ol_m.group(1))
            i += 1
            continue

        # ── Blockquote
        if stripped.startswith(">"):
            text = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def _parse_table_row(row: str) -> list[str]:
    row = row.strip().strip("|")
    return [cell.strip() for cell in row.split("|")]


def _add_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return
    headers = _parse_table_row(table_lines[0])
    # Row 1 is the separator (---), skip it
    data_rows = [_parse_table_row(r) for r in table_lines[2:]]

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    hdr = table.rows[0]
    for j, cell_text in enumerate(headers):
        c = hdr.cells[j]
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(cell_text)
        run.bold = True
        run.font.size = Pt(10)
        # light blue header bg
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)

    # Data rows
    for ri, row_cells in enumerate(data_rows):
        row = table.rows[ri + 1]
        for j, cell_text in enumerate(row_cells):
            if j >= len(row.cells):
                break
            c = row.cells[j]
            c.paragraphs[0].clear()
            add_inline(c.paragraphs[0], cell_text)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


# ── Entry point ───────────────────────────────────────────────────────────────

def main() -> None:
    doc = Document()
    set_doc_styles(doc)

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    parse_md(doc, lines)

    doc.save(OUT_PATH)
    print(f"Gerado: {OUT_PATH}")


if __name__ == "__main__":
    main()
